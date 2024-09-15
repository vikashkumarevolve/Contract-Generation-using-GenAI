from docx import Document
import streamlit as st
import openai
import os
from openai import AzureOpenAI

# Setting up the Azure OpenAI environment according to organization guidelines
openai.api_base = ""
openai.api_type = "azure"
openai.api_version = "2024-02-01"

# Function to get the secure API key
def get_nlp_api_key():
    return ""

# Set the API key using the secure method
openai.api_key = get_nlp_api_key()

# Create the AzureOpenAI client
client = AzureOpenAI(
    azure_endpoint=openai.api_base,
    api_key=openai.api_key,
    api_version=openai.api_version
)

# Contract templates (NDA, Employment, Partnership, Rent)
nda_template = """
THIS NON-DISCLOSURE AGREEMENT (the "Agreement") is entered into on {date} between {company_name}, having its principal place of business at {company_address}, and {partner_name}, having its principal place of business at {partner_address}.
WHEREAS, the parties wish to engage in discussions regarding {purpose_of_agreement};
NOW, THEREFORE, the parties agree to the following terms:
1. Confidentiality: Both parties agree not to disclose confidential information.
2. Term: This Agreement shall remain in effect for {term_of_agreement} years.
...
"""

employment_template = """
THIS EMPLOYMENT AGREEMENT is made as of {date} between {company_name}, a corporation having its principal place of business at {company_address}, and {employee_name}, residing at {employee_address}.
WHEREAS, {company_name} desires to employ {employee_name} as {position_title}, and {employee_name} agrees to accept such employment;
NOW, THEREFORE, the parties agree as follows:
1. Position: {employee_name} will serve as {position_title}.
2. Compensation: {employee_name} shall receive a salary of {salary}.
3. Termination: This agreement may be terminated by either party with {termination_notice} notice.
...
"""

partnership_template = """
THIS PARTNERSHIP AGREEMENT (the "Agreement") is made and entered into as of {date}, by and between {partner1_name} and {partner2_name}.
WHEREAS, the parties wish to form a partnership to carry out the business of {business_description};
NOW, THEREFORE, the parties agree to the following:
1. Partnership: The partnership shall operate under the name {partnership_name}.
2. Capital Contributions: Each partner shall contribute {capital_contribution} to the partnership.
3. Profit Sharing: Profits and losses shall be shared in the following manner: {profit_sharing}.
...
"""

rent_template = """
THIS RENT AGREEMENT is made on {date} at {location} BETWEEN {landlord_name}, residing at {landlord_address}, (hereinafter referred to as "LANDLORD") AND {tenant_name}, residing at {tenant_address} (hereinafter referred to as "TENANT").
NOW, THIS DEED FURTHER WITNESSES:
1. The tenancy term is {term} months, starting from {start_date} and ending on {end_date}.
2. The monthly rent for the property is ₹{monthly_rent}, to be paid on or before the 5th day of each month.
3. A security deposit of ₹{security_deposit} has been paid by the tenant to the landlord.
4. Utilities: The tenant will be responsible for paying utility bills including electricity, water, gas, and others.
5. Termination: Either party may terminate this agreement by giving {termination_notice} days notice.
...
"""

# Function to generate a contract based on user inputs and the contract type
def generate_contract(contract_type, **kwargs):
    """
    Function to generate a contract based on user inputs using Azure OpenAI GPT.
    """
    if contract_type == "NDA":
        contract = nda_template.format(
            date=kwargs['date'],
            company_name=kwargs['company_name'],
            company_address=kwargs['company_address'],
            partner_name=kwargs['partner_name'],
            partner_address=kwargs['partner_address'],
            purpose_of_agreement=kwargs['purpose'],
            term_of_agreement=kwargs['term']
        )
    elif contract_type == "Employment Agreement":
        contract = employment_template.format(
            date=kwargs['date'],
            company_name=kwargs['company_name'],
            company_address=kwargs['company_address'],
            employee_name=kwargs['employee_name'],
            employee_address=kwargs['employee_address'],
            position_title=kwargs['position_title'],
            salary=kwargs['salary'],
            termination_notice=kwargs['termination_notice']
        )
    elif contract_type == "Partnership Agreement":
        contract = partnership_template.format(
            date=kwargs['date'],
            partner1_name=kwargs['partner1_name'],
            partner2_name=kwargs['partner2_name'],
            business_description=kwargs['business_description'],
            partnership_name=kwargs['partnership_name'],
            capital_contribution=kwargs['capital_contribution'],
            profit_sharing=kwargs['profit_sharing']
        )
    elif contract_type == "Rent Agreement":
        contract = rent_template.format(
            date=kwargs['date'],
            location=kwargs['location'],
            landlord_name=kwargs['landlord_name'],
            landlord_address=kwargs['landlord_address'],
            tenant_name=kwargs['tenant_name'],
            tenant_address=kwargs['tenant_address'],
            term=kwargs['term'],
            start_date=kwargs['start_date'],
            end_date=kwargs['end_date'],
            monthly_rent=kwargs['monthly_rent'],
            security_deposit=kwargs['security_deposit'],
            termination_notice=kwargs['termination_notice']
        )
    else:
        return "Invalid contract type"
    
    # Send the contract to Azure GPT-4 for refinement
    try:
        response = client.chat.completions.create(
            model="gpt-4o",  # Replace with your actual Azure OpenAI GPT model
            messages=[{"role": "user", "content": f"Please refine this {contract_type} contract:\n\n{contract}"}],
            max_tokens=500
        )
        
        # Access the message content using dot notation
        if response.choices and hasattr(response.choices[0], "message"):
            generated_contract = response.choices[0].message.content.strip()
            return generated_contract
        else:
            return "No content in the response"

    except Exception as e:
        return f"Error generating contract: {e}"

# Word (.docx) Export function with specific save path
def export_to_word(contract_text, filename="draft_test.docx"):
    save_path = ""  # Specify the folder path
    
    # Ensure the directory exists, if not, create it
    if not os.path.exists(save_path):
        try:
            os.makedirs(save_path)
        except Exception as e:
            st.error(f"Error creating directory: {e}")
            return
    
    full_path = os.path.join(save_path, filename)
    
    # Try saving the document and handle any errors
    try:
        doc = Document()
        doc.add_heading('Contract Document', 0)
        doc.add_paragraph(contract_text)
        doc.save(full_path)
        st.success(f"Contract saved successfully at {full_path}")
    except Exception as e:
        st.error(f"Error saving document: {e}")

# Streamlit UI
st.title("Contract Generator with DOCX Export")

# Select contract type
contract_type = st.selectbox("Select Contract Type", ["NDA", "Employment Agreement", "Partnership Agreement", "Rent Agreement"])

# Shared inputs
date = st.text_input("Date", "Today's Date")

if contract_type == "NDA":
    company_name = st.text_input("Company Name", "ABC Corp")
    company_address = st.text_input("Company Address", "123 Main St")
    partner_name = st.text_input("Partner Name", "XYZ Solutions")
    partner_address = st.text_input("Partner Address", "456 Partner Rd")
    purpose = st.text_input("Purpose of Agreement", "Discussing partnership opportunities")
    term = st.number_input("Term of Agreement (Years)", min_value=1, max_value=10, value=3)

elif contract_type == "Employment Agreement":
    company_name = st.text_input("Company Name", "ABC Corp")
    company_address = st.text_input("Company Address", "123 Main St")
    employee_name = st.text_input("Employee Name", "John Doe")
    employee_address = st.text_input("Employee Address", "789 Employee St")
    position_title = st.text_input("Position Title", "Software Engineer")
    salary = st.text_input("Salary", "$100,000")
    termination_notice = st.text_input("Termination Notice Period", "30 days")

elif contract_type == "Partnership Agreement":
    partner1_name = st.text_input("Partner 1 Name", "Partner A")
    partner2_name = st.text_input("Partner 2 Name", "Partner B")
    business_description = st.text_input("Business Description", "Software development services")
    partnership_name = st.text_input("Partnership Name", "Tech Partners LLC")
    capital_contribution = st.text_input("Capital Contribution", "$50,000 each")
    profit_sharing = st.text_input("Profit Sharing Arrangement", "50/50")

elif contract_type == "Rent Agreement":
    location = st.text_input("Location", "Ahmedabad, Gujarat")
    landlord_name = st.text_input("Landlord Name", "Vasudev Kanubhai Gothi")
    landlord_address = st.text_input("Landlord Address", "7, Shivam Villa, Opp Aanal Tower, Ahmedabad")
    tenant_name = st.text_input("Tenant Name", "Dhaval Vasudev Gothi")
    tenant_address = st.text_input("Tenant Address", "E-1803, Mantri Manyata Lithos, Bengaluru")
    term = st.number_input("Term of Tenancy (Months)", min_value=1, max_value=24, value=11)
    start_date = st.text_input("Start Date", "25/10/2023")
    end_date = st.text_input("End Date", "25/09/2024")
    monthly_rent = st.text_input("Monthly Rent", "₹10,000")
    security_deposit = st.text_input("Security Deposit", "₹20,000")
    termination_notice = st.text_input("Termination Notice (Days)", "30 days")

# Button to generate contract
if st.button("Generate Contract"):
    if contract_type == "NDA":
        contract = generate_contract(
            contract_type,
            date=date,
            company_name=company_name,
            company_address=company_address,
            partner_name=partner_name,
            partner_address=partner_address,
            purpose=purpose,
            term=term
        )
    elif contract_type == "Employment Agreement":
        contract = generate_contract(
            contract_type,
            date=date,
            company_name=company_name,
            company_address=company_address,
            employee_name=employee_name,
            employee_address=employee_address,
            position_title=position_title,
            salary=salary,
            termination_notice=termination_notice
        )
    elif contract_type == "Partnership Agreement":
        contract = generate_contract(
            contract_type,
            date=date,
            partner1_name=partner1_name,
            partner2_name=partner2_name,
            business_description=business_description,
            partnership_name=partnership_name,
            capital_contribution=capital_contribution,
            profit_sharing=profit_sharing
        )
    elif contract_type == "Rent Agreement":
        contract = generate_contract(
            contract_type,
            date=date,
            location=location,
            landlord_name=landlord_name,
            landlord_address=landlord_address,
            tenant_name=tenant_name,
            tenant_address=tenant_address,
            term=term,
            start_date=start_date,
            end_date=end_date,
            monthly_rent=monthly_rent,
            security_deposit=security_deposit,
            termination_notice=termination_notice
        )
    
    # Display the generated contract
    if contract:
        st.subheader("Generated Contract")
        st.text_area("Contract", value=contract, height=400)

        # Button to export the contract as DOCX
        if st.button("Download Contract as DOCX"):
            export_to_word(contract, "draft_test.docx")
